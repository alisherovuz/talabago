from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import uuid
import re
from datetime import datetime


class DocxService:
    def __init__(self):
        self.output_dir = "/tmp/talabago_docs"
        os.makedirs(self.output_dir, exist_ok=True)
    
    async def create_document(self, content: str, topic: str, paper_type: str) -> str:
        """Create a professional university-level DOCX document."""
        
        doc = Document()
        
        # Setup document styles and margins
        self._setup_document(doc)
        
        # Add title page
        self._add_title_page(doc, topic, paper_type)
        
        # Add table of contents placeholder
        self._add_toc_placeholder(doc)
        
        # Add main content
        self._add_content(doc, content)
        
        # Save document
        filename = f"{paper_type}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _setup_document(self, doc: Document):
        """Setup document margins and default styles."""
        
        # Page margins (GOST standard)
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)
            section.page_height = Cm(29.7)  # A4
            section.page_width = Cm(21)
        
        # Setup styles
        styles = doc.styles
        
        # Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(14)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal_style.paragraph_format.space_after = Pt(0)
        normal_style.paragraph_format.first_line_indent = Cm(1.25)
        
        # Heading 1
        h1_style = styles['Heading 1']
        h1_style.font.name = 'Times New Roman'
        h1_style.font.size = Pt(16)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 0, 0)
        h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1_style.paragraph_format.space_before = Pt(24)
        h1_style.paragraph_format.space_after = Pt(12)
        h1_style.paragraph_format.first_line_indent = Cm(0)
        
        # Heading 2
        h2_style = styles['Heading 2']
        h2_style.font.name = 'Times New Roman'
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(0, 0, 0)
        h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2_style.paragraph_format.space_before = Pt(18)
        h2_style.paragraph_format.space_after = Pt(6)
        h2_style.paragraph_format.first_line_indent = Cm(1.25)
        
        # Heading 3
        h3_style = styles['Heading 3']
        h3_style.font.name = 'Times New Roman'
        h3_style.font.size = Pt(14)
        h3_style.font.bold = True
        h3_style.font.italic = True
        h3_style.font.color.rgb = RGBColor(0, 0, 0)
        h3_style.paragraph_format.space_before = Pt(12)
        h3_style.paragraph_format.space_after = Pt(6)
        h3_style.paragraph_format.first_line_indent = Cm(1.25)
    
    def _add_title_page(self, doc: Document, topic: str, paper_type: str):
        """Add professional title page."""
        
        type_names = {
            "referat": "REFERAT",
            "kurs": "KURS ISHI",
            "diplom": "BITIRUV MALAKAVIY ISHI",
            "prezentatsiya": "PREZENTATSIYA"
        }
        
        # Ministry header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("O'ZBEKISTON RESPUBLIKASI")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        # University name placeholder
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("___________________________________ UNIVERSITETI")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        # Faculty placeholder
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("___________________________________ fakulteti")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        # Large spacing
        for _ in range(4):
            doc.add_paragraph()
        
        # Paper type
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(type_names.get(paper_type, "REFERAT"))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18)
        run.font.bold = True
        
        # Spacing
        doc.add_paragraph()
        
        # Topic
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Mavzu: ")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run = p.add_run(topic)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Large spacing
        for _ in range(6):
            doc.add_paragraph()
        
        # Author info (right aligned)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Bajardi: _________________")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Guruh: _________________")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Tekshirdi: _________________")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        # Large spacing to bottom
        for _ in range(6):
            doc.add_paragraph()
        
        # City and year
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Toshkent — {datetime.now().year}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        # Page break
        doc.add_page_break()
    
    def _add_toc_placeholder(self, doc: Document):
        """Add table of contents placeholder."""
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("MUNDARIJA")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        
        doc.add_paragraph()
        
        # TOC instruction
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[Mundarija avtomatik tuziladi]")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_page_break()
    
    def _add_content(self, doc: Document, content: str):
        """Add main content with proper formatting."""
        
        # Ensure content is string
        if isinstance(content, dict):
            content = content.get('text', str(content))
        content = str(content)
        
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Heading 1 (# or BOB, KIRISH, XULOSA, etc.)
            if line.startswith('# '):
                heading_text = line[2:].strip()
                p = doc.add_heading(heading_text, level=1)
                self._format_heading(p, level=1)
            
            # Heading 2 (##)
            elif line.startswith('## '):
                heading_text = line[3:].strip()
                p = doc.add_heading(heading_text, level=2)
                self._format_heading(p, level=2)
            
            # Heading 3 (###)
            elif line.startswith('### '):
                heading_text = line[4:].strip()
                p = doc.add_heading(heading_text, level=3)
                self._format_heading(p, level=3)
            
            # Bullet list
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, text)
            
            # Numbered list
            elif re.match(r'^\d+[\.\)]\s', line):
                match = re.match(r'^(\d+[\.\)])\s*(.*)$', line)
                if match:
                    p = doc.add_paragraph(style='List Number')
                    self._add_formatted_text(p, match.group(2))
            
            # Bibliography entry (starts with number and dot, has author/book pattern)
            elif re.match(r'^\d+\.\s+[A-ZА-ЯЎҚҒҲa-zа-яўқғҳ]', line):
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(1.25)
                p.paragraph_format.hanging_indent = Cm(-1.25)
                self._add_formatted_text(p, line)
            
            # Regular paragraph
            else:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._add_formatted_text(p, line)
    
    def _format_heading(self, paragraph, level: int):
        """Apply consistent heading formatting."""
        
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            if level == 1:
                run.font.size = Pt(16)
                run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                run.font.size = Pt(14)
                run.font.bold = True
            elif level == 3:
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.italic = True
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with bold/italic formatting support."""
        
        # Process **bold** and *italic* markers
        pattern = r'(\*\*.*?\*\*|\*.*?\*)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = paragraph.add_run(part[1:-1])
                run.font.italic = True
            else:
                # Normal text
                run = paragraph.add_run(part)
            
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)


docx_service = DocxService()
